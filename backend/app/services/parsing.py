from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.config import get_settings
from app.services.chunking import ChunkItem, chunk_text

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}


@dataclass
class ParsedDocument:
    chunks: list[ChunkItem]
    raw_text: str


def parse_document(path: str | Path) -> ParsedDocument:
    path = Path(path)
    extension = path.suffix.lower()
    if extension == ".pdf":
        return _parse_pdf(path)
    if extension == ".docx":
        return _parse_docx(path)
    if extension in {".md", ".markdown", ".txt"}:
        return _parse_text(path)
    raise ValueError(f"不支持的文件类型: {extension}")


def _parse_pdf(path: Path) -> ParsedDocument:
    settings = get_settings()
    reader = PdfReader(str(path))
    items: list[ChunkItem] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for part in chunk_text(text, settings.chunk_size, settings.chunk_overlap):
            items.append(ChunkItem(content=part, page_number=page_index))
    raw_text = "\n\n".join(item.content for item in items)
    return ParsedDocument(chunks=items, raw_text=raw_text)


def _parse_docx(path: Path) -> ParsedDocument:
    settings = get_settings()
    doc = DocxDocument(str(path))
    text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    items = [ChunkItem(content=part) for part in chunk_text(text, settings.chunk_size, settings.chunk_overlap)]
    return ParsedDocument(chunks=items, raw_text=text)


def _parse_text(path: Path) -> ParsedDocument:
    settings = get_settings()
    text = path.read_text(encoding="utf-8", errors="ignore")
    items = [ChunkItem(content=part) for part in chunk_text(text, settings.chunk_size, settings.chunk_overlap)]
    return ParsedDocument(chunks=items, raw_text=text)
